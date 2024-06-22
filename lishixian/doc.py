"""Document Operation"""

import os
import csv


def num2str(cell):
    if isinstance(cell, float) and cell.is_integer():
        cell = int(cell)
    return str(cell)


def str2value(s):
    try:
        v = float(s)
        if v.is_integer():
            v = int(v)
    except ValueError:
        v = {'none': None, 'null': None, '': None,
             'true': True, 'flase': False
        }.get(s.lower(), s)
    return v


__all__ = list(globals())


# ---------------------------------------------------------------------------
# Text
# ---------------------------------------------------------------------------


readb = lambda path: (lambda f: [f.read(), f.close()][0])(open(path, 'rb'))


def read(path, encoding='u8', strict=True):
    try:
        if not encoding:
            with open(path, 'rb') as f:
                return f.read()
        for en in [encoding, 'u8', 'u16', 'gbk', None]:
            try:
                with open(path, encoding=en) as f:
                    return f.read()
            except UnicodeError:
                pass
        with open(path, encoding=encoding) as f:
            return f.read()
    except UserWarning if strict else Exception:
        return '' if encoding else b''


def write(path, data='', encoding='u8'):
    folder = os.path.dirname(path)
    if folder and not os.path.exists(folder):
        os.makedirs(folder)
    if isinstance(data, (bytes, bytearray)):
        with open(path, 'wb') as f:
            f.write(data)
    elif isinstance(data, list):
        with open(path, 'w', encoding=encoding) as f:
            f.write('\n'.join(map(str, data)))
    else:
        with open(path, 'w', encoding=encoding) as f:
            f.write(str(data))


def ReadIni(path, encoding='u8'):
    import configparser
    parser = configparser.ConfigParser()
    parser.optionxform = str  # fix opinion can't read key with upper case
    parser.read(path, encoding=encoding)
    return {section: dict(parser.items(section)) for section in parser.sections()}


def WriteIni(path, dic, encoding='u8'):
    import configparser
    parser = configparser.ConfigParser()
    parser.optionxform = str  # fix opinion can't read key with upper case
    parser.read_dict(dic)
    with open(path, 'w', encoding=encoding) as f:
        parser.write(f, False)


def ReadTxt(path, encoding='u8', sep=None):
    with open(path, 'r', encoding=encoding) as f:
        return [line.split(sep) for line in f.read().splitlines()]


def WriteTxt(path, data, encoding='u8', sep=' '):
    with open(path, 'w', encoding=encoding) as f:
        f.write('\n'.join(sep.join(str(cell) for cell in row) for row in data))


def ReadCsv(path, encoding='utf-8-sig'):
    with open(path, encoding=encoding) as f:
        return list(csv.reader(f))


def WriteCsv(path, data, encoding='utf-8-sig', errors='ignore'):
    with open(path, 'w', newline='', encoding=encoding, errors=errors) as f:
        writer = csv.writer(f)
        writer.writerows(data)


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------


def WriteExcel(path, data, new_sheet='sheet1'):
    import xlwt
    xls = xlwt.Workbook('u8')
    sheet = xls.add_sheet(new_sheet, True)
    for r, row in enumerate(data):
        for c, cell in enumerate(row):
            sheet.write(r, c, cell)
    xls.save(path)


def OpenExcel(path):
    import xlrd
    import xlutils.filter

    rb = xlrd.open_workbook(path, formatting_info=True)

    # 参考xlutils.copy库内的用法 参考xlutils.filter内的参数定义style_list
    w = xlutils.filter.XLWTWriter()
    xlutils.filter.process(xlutils.filter.XLRDReader(rb, 'unknown.xls'), w)
    wb = w.output[0][1]
    wb.style_list = w.style_list
    wb.sheets = rb.sheets()

    # quick `style` reach, and `write` with copied style.
    wb.style = lambda n, r, c: wb.style_list[wb.sheets[n].cell_xf_index(r, c)]
    wb.swrite = lambda n, r, c, value: wb.get_sheet(n).write(r, c, value, wb.style(n, r, c))

    return wb


def MergeCell(data, merge, merge_x=True, merge_y=True, strip_x=False):
    import copy
    data2 = []
    for sheet_data, sheet_merge in zip(data, merge):
        sheet_data2 = copy.deepcopy(sheet_data)
        # merge cell
        for r1, r2, c1, c2 in sheet_merge:
            for r in range(r1, r2):
                for c in range(c1, c2):
                    if (not merge_x and c > c1) or (not merge_y and r > r1):
                        sheet_data2[r][c] = None if strip_x else ''
                    else:
                        sheet_data2[r][c] = sheet_data[r1][c1]
        # strip x
        if strip_x:
            sheet_data2 = [[cell for cell in row if cell is not None] for row in sheet_data2]
        data2.append(sheet_data2)
        # remove blanks in tail
        # for row in sheet_data:
        #     while len(row) and row[-1] == '':  # Good!
        #         row.pop()
    return data2


def ReadExcel(path, merge_x=True, merge_y=True, strip_x=False):
    import xlrd
    try:
        xls = xlrd.open_workbook(path, formatting_info=True)
    except xlrd.biffh.XLRDError:
        xls = xlrd.open_workbook(path)

    data = []
    for sheet in xls.sheets():
        sheet_data = []
        for row in range(sheet.nrows):
            sheet_data.append(list(map(num2str, sheet.row_values(row))))
        data.append(sheet_data)

    # only ".xls" type contain merge_info
    merge = [sorted(sheet.merged_cells) for sheet in xls.sheets()]
    data2 = MergeCell(data, merge, merge_x, merge_y, strip_x)
    return data2, merge


def ReadSheet(path, index=0):
    import xlrd
    xls = xlrd.open_workbook(path)
    sheet = xls.sheet_by_index(index)
    return [list(map(num2str, sheet.row_values(row))) for row in range(sheet.nrows)]


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------


def Doc2Docx(path, overwrite=False):
    root, ext = os.path.splitext(path)
    path2 = root + '.docx.temp'
    if overwrite or not os.path.exists(path2):
        from win32com import client
        word = client.Dispatch('Word.Application')
        doc = word.Documents.Open(path)
        doc.SaveAs(path2, 16)
        doc.Close()
    return path2


def OpenDocx(path):
    import docx
    data = []
    merge = []
    doc = docx.Document(path)
    for table in doc.tables:
        cells  = table._cells  # see usage at `docx.table.Table._cells`
        cols   = table._column_count
        length = len(cells)

        # read text
        data.append([])
        for i, cell in enumerate(cells):
            r, c = divmod(i, cols)
            if c == 0:  # create new row
                data[-1].append([])
            data[-1][-1].append(cell.text)

        # find merged
        merge.append([])
        for i, cell in enumerate(cells):
            if cell in cells[:i]:  # only find first repeated cell
                continue
            for j in range(length - 1, 0, -1):  # find last repeated cell
                if cell is cells[j]:  # always can be found
                    break
            if i != j:  # first != last -> merged cells
                r1, c1 = divmod(i, cols)
                r2, c2 = divmod(j, cols)
                merge[-1].append((r1, r2 + 1, c1, c2 + 1))

    return data, merge


def ReadWordTexts(path):
    import docx
    if path.endswith('.doc'):
        path = Doc2Docx(path)
    doc = docx.Document(path)
    return [p.text for p in doc.paragraphs]


def ReadWord(path, merge_x=True, merge_y=True, strip_x=False):
    if path.endswith('.doc'):
        path = Doc2Docx(path)
    data, merge = OpenDocx(path)
    data2 = MergeCell(data, merge, merge_x, merge_y, strip_x)
    return data2


# ---------------------------------------------------------------------------
# Other
# ---------------------------------------------------------------------------


def ReadFile(path, merge_x=True, merge_y=True, strip_x=False):
    ext = os.path.splitext(path)[1]
    if ext in ['.csv', '.txt']:
        return ReadCsv(path)
    elif ext in ['.xls', '.xlsx']:
        return ReadExcel(path, merge_x, merge_y, strip_x)
    elif ext in ['.doc', '.docx']:
        return ReadWord(path, merge_x, merge_y, strip_x)


def File2Csv(path, merge_x=True, merge_y=True, strip_x=False):
    root, ext = os.path.splitext(path)
    data = ReadFile(path, merge_x, merge_y, strip_x)
    WriteCsv(data, root + '.csv')
    return data


def ReadFiles(paths, merge_x=True, merge_y=True, strip_x=False):
    data = []
    for path in paths:
        folder, filename = os.path.split(path)
        data.extend([[folder, filename] + row for row in ReadFile(path, merge_x, merge_y, strip_x)])
    return data


__all__ = [k for k in globals() if k not in __all__]


if __name__ == '__main__':
    from pprint import pprint

    # test read
    data = ReadExcel('../test.xls')
    # pprint(data)

    pprint(ReadWord('../test.docx', False, False, False))
    pprint(ReadWord('../test.docx', False, True,  False))
    pprint(ReadWord('../test.docx', True,  False, False))
    pprint(ReadWord('../test.docx', True,  True,  False))

    # test read/write with keep style
    wb = OpenExcel('../test2.xls')
    sheet2 = wb.get_sheet(0)
    for r in range(3):
        for c in range(3):
            sheet2.write(r, c, 'T%d%d' % (r, c), wb.style(0, r, c))
            # wb.swrite(0, r, c, 'V%d%d' % (r, c))

    wb.save('save2.xls')
